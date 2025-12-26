"""Text extraction agent - handles PDF, DOCX, emails, and ENEX files."""

import re
import html
from pathlib import Path
from typing import Dict, Any, Optional
import xml.etree.ElementTree as ET
from datetime import datetime

from .base_agent import BaseAgent, DocumentData


class ExtractionAgent(BaseAgent):
    """Agent for extracting text from various document formats."""

    def __init__(self, config: Dict[str, Any], logger=None):
        """Initialize extraction agent.

        Args:
            config: Configuration dictionary
            logger: Optional logger instance
        """
        super().__init__(config, logger)
        self.extraction_config = config.get('extraction', {})

        # Initialize extractors
        self._init_pdf_extractor()
        self._init_docx_extractor()
        self._init_email_extractor()

    def _init_pdf_extractor(self):
        """Initialize PDF extraction libraries."""
        try:
            import fitz  # PyMuPDF
            self.pymupdf = fitz
            self.logger.info("PyMuPDF initialized for PDF extraction")
        except ImportError:
            self.pymupdf = None
            self.logger.warning("PyMuPDF not available")

        try:
            import pdfplumber
            self.pdfplumber = pdfplumber
            self.logger.info("pdfplumber initialized as PDF fallback")
        except ImportError:
            self.pdfplumber = None
            self.logger.warning("pdfplumber not available")

    def _init_docx_extractor(self):
        """Initialize DOCX extraction library."""
        try:
            import docx
            self.docx = docx
            self.logger.info("python-docx initialized for DOCX extraction")
        except ImportError:
            self.docx = None
            self.logger.warning("python-docx not available")

    def _init_email_extractor(self):
        """Initialize email extraction libraries."""
        try:
            import extract_msg
            self.extract_msg = extract_msg
            self.logger.info("extract_msg initialized for .msg files")
        except ImportError:
            self.extract_msg = None
            self.logger.warning("extract_msg not available")

        try:
            from email import message_from_file
            from email.policy import default
            self.email_parser = (message_from_file, default)
            self.logger.info("email.parser initialized for .eml files")
        except ImportError:
            self.email_parser = None
            self.logger.warning("email.parser not available")

    def process(self, doc: DocumentData) -> DocumentData:
        """Extract text from document.

        Args:
            doc: DocumentData object

        Returns:
            DocumentData with extracted text in content field
        """
        try:
            if doc.file_type == 'pdf':
                doc = self._extract_pdf(doc)
            elif doc.file_type in {'docx', 'doc'}:
                doc = self._extract_docx(doc)
            elif doc.file_type == 'enex':
                doc = self._extract_enex(doc)
            elif doc.file_type in {'eml', 'msg'}:
                doc = self._extract_email(doc)
            elif doc.file_type in {'txt', 'md', 'text'}:
                doc = self._extract_text(doc)
            else:
                self.log_skip(f"{doc.file_path.name}: Unsupported format for text extraction")
                return doc

            if doc.content:
                doc.processing_status = 'completed'
                self.log_success(f"Extracted {len(doc.content)} chars from {doc.file_path.name}")
            else:
                self.log_error(f"No text extracted from {doc.file_path.name}")
                doc.errors.append("No text extracted")

        except Exception as e:
            self.log_error(f"Extraction failed for {doc.file_path.name}", e)
            doc.processing_status = 'failed'
            doc.errors.append(f"Extraction error: {str(e)}")

        return doc

    def _extract_pdf(self, doc: DocumentData) -> DocumentData:
        """Extract text from PDF.

        Args:
            doc: DocumentData with PDF file

        Returns:
            DocumentData with extracted text
        """
        pdf_config = self.extraction_config.get('pdf', {})
        prefer_text = pdf_config.get('prefer_text_layer', True)

        text_parts = []
        page_metadata = []

        # Try PyMuPDF first
        if self.pymupdf and prefer_text:
            try:
                pdf = self.pymupdf.open(stream=doc.raw_data, filetype="pdf")

                for page_num, page in enumerate(pdf, start=1):
                    page_text = page.get_text()

                    if page_text.strip():
                        text_parts.append(page_text)
                        page_metadata.append({
                            'page_number': page_num,
                            'char_count': len(page_text)
                        })

                pdf.close()

                # Update metadata
                doc.metadata['num_pages'] = len(pdf)
                doc.metadata['pages'] = page_metadata

                if text_parts:
                    doc.content = '\n\n'.join(text_parts)
                    self.logger.debug(f"PyMuPDF extracted text from {len(text_parts)} pages")
                    return doc

            except Exception as e:
                self.logger.warning(f"PyMuPDF extraction failed: {e}")

        # Fallback to pdfplumber
        if self.pdfplumber:
            try:
                import io
                pdf = self.pdfplumber.open(io.BytesIO(doc.raw_data))

                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text()

                    if page_text:
                        text_parts.append(page_text)
                        page_metadata.append({
                            'page_number': page_num,
                            'char_count': len(page_text)
                        })

                pdf.close()

                doc.metadata['num_pages'] = len(pdf.pages)
                doc.metadata['pages'] = page_metadata

                if text_parts:
                    doc.content = '\n\n'.join(text_parts)
                    self.logger.debug(f"pdfplumber extracted text from {len(text_parts)} pages")
                    return doc

            except Exception as e:
                self.logger.warning(f"pdfplumber extraction failed: {e}")

        # If no text layer found, mark for OCR
        if not text_parts:
            self.logger.info(f"{doc.file_path.name}: No text layer found, requires OCR")
            doc.metadata['requires_ocr'] = True

        return doc

    def _extract_docx(self, doc: DocumentData) -> DocumentData:
        """Extract text from DOCX file.

        Args:
            doc: DocumentData with DOCX file

        Returns:
            DocumentData with extracted text
        """
        if not self.docx:
            raise RuntimeError("python-docx not available")

        import io
        document = self.docx.Document(io.BytesIO(doc.raw_data))

        text_parts = []
        docx_config = self.extraction_config.get('docx', {})

        # Extract paragraphs
        for para in document.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in document.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        # Extract headers/footers if configured
        if docx_config.get('extract_headers_footers', True):
            for section in document.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            text_parts.insert(0, f"[HEADER] {para.text}")

                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            text_parts.append(f"[FOOTER] {para.text}")

        # Metadata
        core_props = document.core_properties
        if core_props.author:
            doc.metadata['author'] = core_props.author
        if core_props.created:
            doc.metadata['created_date'] = core_props.created.isoformat()
        if core_props.modified:
            doc.metadata['modified_date'] = core_props.modified.isoformat()
        if core_props.title:
            doc.metadata['title'] = core_props.title

        doc.content = '\n\n'.join(text_parts)
        return doc

    def _extract_enex(self, doc: DocumentData) -> DocumentData:
        """Extract text from Evernote ENEX file.

        Args:
            doc: DocumentData with ENEX file

        Returns:
            DocumentData with extracted text (combines all notes)
        """
        tree = ET.parse(doc.file_path)
        root = tree.getroot()

        notes = []

        for note in root.findall('.//note'):
            note_data = self._parse_enex_note(note)
            notes.append(note_data['content'])

            # Store first note's metadata as document metadata
            if len(notes) == 1:
                doc.metadata.update({
                    'title': note_data.get('title', ''),
                    'created_date': note_data.get('created', ''),
                    'tags': note_data.get('tags', [])
                })

        doc.metadata['num_notes'] = len(notes)
        doc.content = '\n\n================================================================================\n\n'.join(notes)

        return doc

    def _parse_enex_note(self, note_element: ET.Element) -> Dict[str, Any]:
        """Parse individual note from ENEX.

        Args:
            note_element: XML element for note

        Returns:
            Dictionary with note data
        """
        # Extract metadata
        title = note_element.find('title')
        title = title.text if title is not None else 'Untitled'

        created = note_element.find('created')
        created_str = self._parse_enex_date(created.text) if created is not None else ''

        updated = note_element.find('updated')
        updated_str = self._parse_enex_date(updated.text) if updated is not None else ''

        tags = [tag.text for tag in note_element.findall('.//tag') if tag.text]

        # Extract content
        content_elem = note_element.find('content')
        content = ''
        if content_elem is not None and content_elem.text:
            content = self._extract_enex_content(content_elem.text)

        # Format note
        note_text = f"# {title}\n\n"
        note_text += f"**Created:** {created_str}\n"
        note_text += f"**Updated:** {updated_str}\n"
        if tags:
            note_text += f"**Tags:** {', '.join(tags)}\n"
        note_text += f"\n---\n\n{content}"

        return {
            'title': title,
            'created': created_str,
            'updated': updated_str,
            'tags': tags,
            'content': note_text
        }

    def _extract_enex_content(self, content: str) -> str:
        """Extract text from ENML content.

        Args:
            content: ENML content string

        Returns:
            Cleaned text
        """
        # Remove CDATA
        content = re.sub(r'<!\[CDATA\[(.*?)\]\]>', r'\1', content, flags=re.DOTALL)

        # Convert links
        content = re.sub(r'<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>', r'[\2](\1)', content)

        # Convert line breaks
        content = content.replace('<br/>', '\n').replace('<br>', '\n')
        content = re.sub(r'</p>', '\n\n', content)
        content = re.sub(r'</div>', '\n', content)

        # Remove remaining HTML tags
        content = re.sub(r'<[^>]+>', '', content)

        # Decode HTML entities
        content = html.unescape(content)

        # Clean up whitespace
        content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)
        content = content.strip()

        return content

    def _parse_enex_date(self, date_str: str) -> str:
        """Parse Evernote date format.

        Args:
            date_str: Date in format YYYYMMDDTHHMMSSZ

        Returns:
            ISO formatted date string
        """
        try:
            dt = datetime.strptime(date_str, '%Y%m%dT%H%M%SZ')
            return dt.isoformat()
        except Exception:
            return date_str

    def _extract_email(self, doc: DocumentData) -> DocumentData:
        """Extract text from email files.

        Args:
            doc: DocumentData with email file

        Returns:
            DocumentData with extracted text
        """
        if doc.file_type == 'msg' and self.extract_msg:
            return self._extract_msg(doc)
        elif doc.file_type == 'eml' and self.email_parser:
            return self._extract_eml(doc)
        else:
            raise RuntimeError(f"No email extractor available for {doc.file_type}")

    def _extract_msg(self, doc: DocumentData) -> DocumentData:
        """Extract text from .msg file.

        Args:
            doc: DocumentData with .msg file

        Returns:
            DocumentData with extracted text
        """
        msg = self.extract_msg.Message(doc.file_path)

        doc.metadata.update({
            'subject': msg.subject or '',
            'sender': msg.sender or '',
            'date': msg.date or '',
            'to': msg.to or '',
        })

        # Combine email parts
        parts = []
        parts.append(f"From: {msg.sender}")
        parts.append(f"To: {msg.to}")
        parts.append(f"Subject: {msg.subject}")
        parts.append(f"Date: {msg.date}")
        parts.append("\n---\n")
        parts.append(msg.body or '')

        doc.content = '\n'.join(parts)
        msg.close()

        return doc

    def _extract_eml(self, doc: DocumentData) -> DocumentData:
        """Extract text from .eml file.

        Args:
            doc: DocumentData with .eml file

        Returns:
            DocumentData with extracted text
        """
        message_from_file, default = self.email_parser

        with open(doc.file_path, 'r', encoding='utf-8', errors='ignore') as f:
            msg = message_from_file(f, policy=default)

        doc.metadata.update({
            'subject': msg.get('subject', ''),
            'from': msg.get('from', ''),
            'to': msg.get('to', ''),
            'date': msg.get('date', ''),
        })

        # Extract body
        body = ''
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == 'text/plain':
                    body = part.get_content()
                    break
        else:
            body = msg.get_content()

        # Combine email parts
        parts = []
        parts.append(f"From: {msg.get('from', '')}")
        parts.append(f"To: {msg.get('to', '')}")
        parts.append(f"Subject: {msg.get('subject', '')}")
        parts.append(f"Date: {msg.get('date', '')}")
        parts.append("\n---\n")
        parts.append(body)

        doc.content = '\n'.join(parts)

        return doc

    def _extract_text(self, doc: DocumentData) -> DocumentData:
        """Extract text from plain text files.

        Args:
            doc: DocumentData with text file

        Returns:
            DocumentData with extracted text
        """
        try:
            with open(doc.file_path, 'r', encoding='utf-8', errors='ignore') as f:
                doc.content = f.read()
        except Exception as e:
            # Try to detect encoding
            import chardet
            with open(doc.file_path, 'rb') as f:
                raw = f.read()
                detected = chardet.detect(raw)
                encoding = detected.get('encoding', 'utf-8')

            with open(doc.file_path, 'r', encoding=encoding, errors='ignore') as f:
                doc.content = f.read()

        return doc
